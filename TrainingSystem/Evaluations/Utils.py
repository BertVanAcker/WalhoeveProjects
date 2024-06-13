import pandas
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches
import math
import numpy as np

class EvaluationAnalyzer():
    """
            Module: Class representing an importer module
             :param string Name: Name of the User
             :param string Role: Role of the User
        """

    def __init__(self, Name="",excelSheet=None):
        self.Name = Name
        self.RAWentries = []

        if excelSheet is not None:
            self._inputFile = excelSheet
            self.importEvaluations(excelSheet=self._inputFile)
        else:
            self._inputFile = ""

        self._currentFigure = None


    @property
    def entriesRAW(self):
        return self.RAWentries

    @property
    def entryCount(self):
        return len(self.RAWentries)

    @property
    def trainingList(self):
        trainingList = []
        #determine unique trainings
        for entry in self.RAWentries:
            title = entry["title"]
            if title not in trainingList:
                trainingList.append(title)
        return trainingList

    def importEvaluations(self,excelSheet=None,googleSheet=None):
        entries=[]

        if excelSheet is not None:
            df = pandas.read_excel(excelSheet)

            submission=df['Submission Date'].values
            firstName = df['First Name'].values
            lastName = df['Last Name'].values
            topic = df['Onderwerp en inhoud van de opleiding'].values
            topicRemarks = df['De opleiding voldeed aan mijn verwachtingen'].values
            teacher = df['Lesgever(s)'].values
            teacherRemarks = df['Lesgever(s) 2'].values
            location = df['Locatie'].values
            locationRemarks = df['Locatie 2'].values
            duration = df['Duur van de opleiding'].values
            durationRemarks = df['Duur van de opleiding 2'].values
            improvement = df['Ik kan mijn job nu beter uitvoeren'].values
            improvementRemarks = df['Ik kan mijn job nu beter uitvoeren 2'].values
            expectation = df['De opleiding voldeed aan mijn verwachtingen 2'].values
            expectationRemarks = df['De opleiding voldeed aan mijn verwachtingen 3'].values
            title = df['Titel opleiding'].values

        #generate entries
        for i in range(0,len(submission),1):
            entries.append({"Date":submission[i].__str__(),"title":title[i],"firstName":firstName[i],"lastName":lastName[i],"topicScore":topic[i],"teacherScore":teacher[i],"locationScore":location[i],"durationScore":duration[i],"improvementScore":improvement[i],"expectationScore":expectation[i],"topicRemarks":topicRemarks[i],"teacherRemarks":teacherRemarks[i],"locationRemarks":locationRemarks[i],"durationRemarks":durationRemarks[i],"improvementRemarks":improvementRemarks[i],"expectationRemarks":expectationRemarks[i]})

        self.RAWentries = entries
        return entries

    def _createRadarChart(self, title="", inputData=None, visualize=False, export=True):

        df = pandas.DataFrame(dict(
            r=[inputData["topicScore"], inputData["teacherScore"], inputData["locationScore"], inputData["improvementScore"],
               inputData["expectationScore"]],
            theta=['Onderwerp', 'Lesgever(s)', 'Locatie','Bruikbaarheid', 'Verwachtingen']))
        fig = px.line_polar(df, r='r', theta='theta', line_close=True)
        fig.update_traces(fill='toself')
        if export:
            fig.write_image("output/static/" + title.replace(" ", "") + "_spiderchart.png")
            fig.write_html("output/dynamic/" + title.replace(" ", "") + "_spiderchart.html")
        if visualize:
            fig.show()

    def _createBarChart(self, title="", inputData=None, visualize=False, export=True):

        scoreOptions = ['0(Slecht)', '1', '2', '3', '4', '5(Heel goed)']

        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=scoreOptions,
            y=inputData["topicScore"],
            name='Onderwerp',
            marker_color='indianred'
        ))
        fig.add_trace(go.Bar(
            x=scoreOptions,
            y=inputData["teacherScore"],
            name='Lesgever(s)',
            marker_color='lightsalmon'
        ))
        fig.add_trace(go.Bar(
            x=scoreOptions,
            y=inputData["locationScore"],
            name='Locatie',
            marker_color='limegreen'
        ))
        fig.add_trace(go.Bar(
            x=scoreOptions,
            y=inputData["improvementScore"],
            name='Bruikbaarheid',
            marker_color='darkviolet'
        ))
        fig.add_trace(go.Bar(
            x=scoreOptions,
            y=inputData["expectationScore"],
            name='Verwachting',
            marker_color='darkturquoise'
        ))

        if export:
            fig.write_image("output/static/" + title.replace(" ", "") + "_barchart.png")
            fig.write_html("output/dynamic/" + title.replace(" ", "") + "_barchart.html")
        if visualize:
            fig.show()
    def compareTrainings(self,list=[],visualize=False,export=True):

        chartInput = []
        for training in list:
            count = 0
            scores = {'topicScore': 0.0, 'teacherScore': 0.0, 'locationScore': 0.0, 'durationScore': 0.0,
                      'improvementScore': 0.0, 'expectationScore': 0.0}
            # loop over all entries and calculate mean()
            for entry in self.RAWentries:
                if entry["title"] == training:
                    count = count + 1
                    scores["topicScore"] = scores["topicScore"] + entry["topicScore"]
                    scores["teacherScore"] = scores["teacherScore"] + entry["teacherScore"]
                    scores["locationScore"] = scores["locationScore"] + entry["locationScore"]
                    scores["improvementScore"] = scores["improvementScore"] + entry["improvementScore"]
                    scores["expectationScore"] = scores["expectationScore"] + entry["expectationScore"]
            # post-process mean
            scores["topicScore"] = scores["topicScore"] / count
            scores["teacherScore"] = scores["teacherScore"] / count
            scores["locationScore"] = scores["locationScore"] / count
            scores["improvementScore"] = scores["improvementScore"] / count
            scores["expectationScore"] = scores["expectationScore"] / count

            chartInput.append(scores)

        categories = ['Onderwerp', 'Lesgever(s)', 'Locatie','Bruikbaarheid', 'Verwachtingen']

        fig = go.Figure()
        for i in range(0,len(list),1):
            fig.add_trace(go.Scatterpolar(
                r=[chartInput[i]["topicScore"], chartInput[i]["teacherScore"], chartInput[i]["locationScore"], chartInput[i]["improvementScore"], chartInput[i]["expectationScore"]],
                theta=categories,
                fill='toself',
                name=list[i]
            ))
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 5]
                )),
            showlegend=True
        )
        if export:
            fig.write_image("output/comparison/comparison.png")
            fig.write_html("output/comparison/comparison.html")

        if visualize:
            fig.show()
        self._currentFigure = fig

    def generateRadarChart(self,title=None,visualize=False,export=True):
        _uniqueTrainingNames = self.trainingList

        if title is None:
            #TODO: generate the radar charts of all trainings and map them in a single chart
            x = 10
        elif title == "ALL":
            # TODO: generate the radar charts of all trainings in a seperate chart
            for training in _uniqueTrainingNames:
                count = 0
                scores = {'topicScore': 0.0, 'teacherScore': 0.0, 'locationScore': 0.0, 'durationScore': 0.0,
                          'improvementScore': 0.0, 'expectationScore': 0.0}
                # loop over all entries and calculate mean()
                for entry in self.RAWentries:
                    if entry["title"] == training:
                        count = count + 1
                        scores["topicScore"] = scores["topicScore"] + entry["topicScore"]
                        scores["teacherScore"] = scores["teacherScore"] + entry["teacherScore"]
                        scores["locationScore"] = scores["locationScore"] + entry["locationScore"]
                        scores["improvementScore"] = scores["improvementScore"] + entry["improvementScore"]
                        scores["expectationScore"] = scores["expectationScore"] + entry["expectationScore"]
                # post-process mean
                scores["topicScore"] = scores["topicScore"] / count
                scores["teacherScore"] = scores["teacherScore"] / count
                scores["locationScore"] = scores["locationScore"] / count
                scores["improvementScore"] = scores["improvementScore"] / count
                scores["expectationScore"] = scores["expectationScore"] / count

                self._createRadarChart(title=training, inputData=scores, visualize=False, export=export)
        else:
            count=0
            scores = {'topicScore': 0.0, 'teacherScore': 0.0, 'locationScore': 0.0, 'durationScore': 0.0,
                      'improvementScore': 0.0, 'expectationScore': 0.0}
            # loop over all entries and calculate mean()
            for entry in self.RAWentries:
                if entry["title"]== title:
                    count = count + 1
                    scores["topicScore"] = scores["topicScore"] + entry["topicScore"]
                    scores["teacherScore"] = scores["teacherScore"] + entry["teacherScore"]
                    scores["locationScore"] = scores["locationScore"] + entry["locationScore"]
                    scores["improvementScore"] = scores["improvementScore"] + entry["improvementScore"]
                    scores["expectationScore"] = scores["expectationScore"] + entry["expectationScore"]
            #post-process mean
            scores["topicScore"] = scores["topicScore"] / count
            scores["teacherScore"] = scores["teacherScore"] / count
            scores["locationScore"] = scores["locationScore"] / count
            scores["improvementScore"] = scores["improvementScore"] / count
            scores["expectationScore"] = scores["expectationScore"] / count

            self._createRadarChart(title=title, inputData=scores, visualize=visualize, export=export)

    def generateBarChart(self,title=None,visualize=False,export=True):
        _uniqueTrainingNames = self.trainingList

        if title is None:
            #TODO: generate the radar charts of all trainings and map them in a single chart
            x = 10
        elif title == "ALL":
            # TODO: generate the radar charts of all trainings in a seperate chart
            for training in _uniqueTrainingNames:
                scores = {'topicScore': [0,0,0,0,0,0], 'teacherScore': [0,0,0,0,0,0], 'locationScore': [0,0,0,0,0,0], 'durationScore': [0,0,0,0,0,0],
                          'improvementScore': [0,0,0,0,0,0], 'expectationScore': [0,0,0,0,0,0]}

                # loop over all entries and check entries for each category mean()
                for entry in self.RAWentries:
                    if entry["title"] == training:

                        scores["topicScore"][int(entry["topicScore"])] = scores["topicScore"][int(entry["topicScore"])] + 1
                        scores["teacherScore"][int(entry["teacherScore"])] = scores["teacherScore"][int(entry["teacherScore"])] + 1
                        scores["locationScore"][int(entry["locationScore"])] = scores["locationScore"][int(entry["locationScore"])] + 1
                        scores["durationScore"][int(entry["durationScore"])] = scores["durationScore"][int(entry["durationScore"])] + 1
                        scores["improvementScore"][int(entry["improvementScore"])] = scores["improvementScore"][int(entry["improvementScore"])] + 1
                        scores["expectationScore"][int(entry["expectationScore"])] = scores["expectationScore"][int(entry["expectationScore"])] + 1

                self._createBarChart(title=training, inputData=scores, visualize=False, export=export)
        else:
            count=0


    def generateDocuments(self):

        #generate word document per evaluation
        for training in self.trainingList:
            document = Document()

            document.add_heading(training+" overzicht", 0)

            # INSERT SPIDER CHART
            document.add_heading("Scores per categorie", 1)
            document.add_picture('output/static/'+training.replace(" ", "")+'_spiderchart.png',width=Inches(6))

            # INSERT BAR CHART
            document.add_heading("Aantal stemmen per score", 1)
            document.add_picture('output/static/' + training.replace(" ", "") + '_barchart.png', width=Inches(6))

            # INSERT REMARKS
            document.add_heading("Opmerkingen", 1)
            document.add_heading("Onderwerp", 2)
            for entry in self.RAWentries:
                if entry["title"] == training:
                    if not self.isNaN(entry["topicRemarks"]):
                        document.add_paragraph(entry["topicRemarks"], style='List Bullet')

            document.add_heading("Lesgever(s)", 2)
            for entry in self.RAWentries:
                if entry["title"] == training:
                    if not self.isNaN(entry["teacherRemarks"]):
                        document.add_paragraph(entry["teacherRemarks"], style='List Bullet')

            document.add_heading("Locatie", 2)
            for entry in self.RAWentries:
                if entry["title"] == training:
                    if not self.isNaN(entry["locationRemarks"]):
                        document.add_paragraph(entry["locationRemarks"], style='List Bullet')

            document.add_heading("Bruikbaar", 2)
            for entry in self.RAWentries:
                if entry["title"] == training:
                    if not self.isNaN(entry["improvementRemarks"]):
                        document.add_paragraph(entry["improvementRemarks"], style='List Bullet')

            document.add_heading("Verwachting", 2)
            for entry in self.RAWentries:
                if entry["title"] == training:
                    if not self.isNaN(entry["expectationRemarks"]):
                        document.add_paragraph(entry["expectationRemarks"], style='List Bullet')

            document.save('output/documents/'+training.replace(" ", "")+'.docx')

    def isNaN(self,num):
        return num != num


class RiskAnalyzer():
    """
            Module: Class representing an importer module
             :param string Name: Name of the User
             :param string Role: Role of the User
        """

    def __init__(self, Name="",excelSheet=None):
        self.Name = Name
        self.RAWentries = []
        self.ANALYZEDentries = []
        self.RAWages = []
        self.ANALYZEDages = [0, 0, 0]

        if excelSheet is not None:
            self._inputFile = excelSheet
            self._importRiskAnalysis(excelSheet=self._inputFile)
        else:
            self._inputFile = ""

        self.answerOptions = ['Niet van toepassing','Helemaal niet akkoord','Niet akkoord','Akkoord','Helemaal akkoord']
        self.ageOptions = ['18 tem 30 jaar', '31 tem 44 jaar', '45+']


    def _importRiskAnalysis(self,excelSheet=None,googleSheet=None):
        questionList = []
        if excelSheet is not None:
            df = pandas.read_excel(excelSheet)
            # FETCH QUESTIONS
            for i in range(2,df.columns.values.shape[0]-1,1):
                question = []
                question.append(df.columns.values[i])
                for j in range(0,df.values.shape[0],1):
                    question.append(df.values[j][i])

                questionList.append(question)

            # FETCH AGES
            for i in range(0,df.values.shape[0],1):
                self.RAWages.append(df.values[i][1])

        self.RAWentries = questionList
        return questionList

    def _analyzeQuestions(self):

        for answer in self.RAWentries:
            # DETERMINE PERCENTAGE PER ANSWER
            answerCount = [0, 0, 0, 0, 0]
            for i in range(1,len(answer),1):
                for j in range(0,len(self.answerOptions),1):
                    if answer[i] == self.answerOptions[j]:
                        answerCount[j] = answerCount[j]+1

            #name = answer[0].split(") ")[1]
            self.ANALYZEDentries.append([answer[0],answerCount])


        for i in range(0,len(self.RAWages),1):
            for j in range(0, len(self.ageOptions), 1):
                if self.RAWages[i] == self.ageOptions[j]:
                    self.ANALYZEDages[j] = self.ANALYZEDages[j] + 1


    def func(self,pct, allvalues):
        absolute = int(pct / 100. * np.sum(allvalues))
        return "{:.1f}%\n({:d} g)".format(pct, absolute)

    def _createPieChart(self,title=None,inputData=None,labels=None,visualize=False,export=True):

        #fig = px.pie(values=inputData, names=self.answerOptions,sort=True)
        fig = go.Figure(data=go.Pie(values=inputData,labels=labels,direction ='clockwise',sort=False,hole = 0.3,))
        fig.update_traces(marker=dict(colors=["#D37676","#EBC49F","#F1EF99","#B0C5A4","#A4C1C5"]),texttemplate="%{percent} <br>(%{value})" )

        if export:
            fig.write_image("output/static/" + title.replace(" ", "") + "_piechart.png")
        if visualize:
            fig.show()
    def createAllPieCharts(self):

        for i in range(0,len(self.ANALYZEDentries),1):
            self._createPieChart(title="Question_"+(i+1).__str__(), visualize=False,labels=self.answerOptions, export=True, inputData=self.ANALYZEDentries[i][1])

        #generate age chart
        self._createPieChart(title="Ages", visualize=False, labels=self.ageOptions,export=True, inputData=self.ANALYZEDages)

    def generateDocument(self):


        #generate word document for the risk analysis
        document = Document()

        document.add_heading("Resultaten psychosociale risicoanalyse", 0)

        document.add_heading("1) Leeftijdscategorieën", 1)
        document.add_picture('output/static/Ages_piechart.png', width=Inches(6))

        for i in range(0,len(self.ANALYZEDentries),1):
            document.add_heading(self.ANALYZEDentries[i][0], 1)
            document.add_picture('output/static/Question_'+(i+1).__str__()+'_piechart.png',width=Inches(6))



            document.save('output/ResultatenPsychosocialeRisicoanalyse.docx')
